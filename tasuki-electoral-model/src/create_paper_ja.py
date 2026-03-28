#!/usr/bin/env python3
"""Generate Japanese JASSS paper as .docx with embedded color figures."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

FIGS = '/home/ubuntu/figures'
OUT = '/home/ubuntu/TASUKI_JASSS_Japanese.docx'

doc = Document()

# ── Style Setup ──────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
# Set East Asian font
rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
if rFonts is None:
    from docx.oxml.ns import qn
    from lxml import etree
    rPr = style.element.get_or_add_rPr()
    rFonts_elem = rPr.find(qn('w:rFonts'))
    if rFonts_elem is None:
        rFonts_elem = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts_elem.set(qn('w:eastAsia'), 'Yu Mincho')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Arial'
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    if level == 1:
        hs.font.size = Pt(16)
    elif level == 2:
        hs.font.size = Pt(13)
    else:
        hs.font.size = Pt(11)

def add_para(text, bold=False, italic=False, align=None, size=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_figure(path, caption, width=5.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)
    cap.paragraph_format.space_after = Pt(12)

# ══════════════════════════════════════════════
# 表紙
# ══════════════════════════════════════════════
add_para('', space_after=24)
add_para(
    '統合的知識に基づく信頼調整型評価システム（TASUKI 襟）：\n'
    '説明責任駆動型選挙制度改革のエージェントベースモデル',
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18, space_after=12
)
add_para(
    'Trust-Adjusted Scoring with Unified Knowledge Integration (TASUKI):\n'
    'An Agent-Based Model of Accountability-Driven Electoral Reform',
    italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=12
)
add_para('[査読用匿名版 \u2014 著者名・所属機関非開示]',
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=24)

# ── 要旨 ──
doc.add_heading('要旨（Abstract）', level=1)
doc.add_paragraph(
    '民主主義的選挙は代議士の説明責任を確保する主要なメカニズムであるが、選挙的制裁の二値的性質'
    '（再選か否か）は公約実現のインセンティブとして粗い手段に留まる。本論文では、候補者信任係数を'
    '用いた動的実績連動型代議制（TASUKI: Dynamic Retrospective Delegation with Candidate Trust）'
    'を提案する。TASUKIでは、候補者が重み付き公約ポートフォリオを事前宣言し、任期終了時に第三者'
    '評価機関が実現度を査定し、その結果得られる説明責任得点を影響関数によって候補者レベルの信任'
    '係数に変換して次回選挙の有効得票数に反映する。ODDプロトコルに準拠したエージェントベースモデル'
    'により多世代選挙動態をシミュレーションした結果、TASUKIは(i)標準選挙と比較して平均説明責任'
    '得点を有意に向上させ、(ii)誠実な候補者を有利とする進化的淘汰圧を生じさせ、(iii)遺伝的'
    'アルゴリズム探索により発見された多様な敵対的戦略に対してロバスト性を示すことが確認された。'
    '感度分析の結果、凹関数およびシグモイド型影響関数がインセンティブ強度と操作耐性の最適な'
    'トレードオフを提供することが明らかになった。'
)

add_para('キーワード：選挙的説明責任、エージェントベースモデル、メカニズムデザイン、'
         '回顧的投票、加重信任、公約実現度、ODDプロトコル',
         italic=True, size=10, space_after=18)

# ══════════════════════════════════════════════
# 1. 序論
# ══════════════════════════════════════════════
doc.add_heading('1. 序論（Introduction）', level=1)

doc.add_paragraph(
    '選挙は代議制民主主義の根幹であり、市民が代議士を定期的に評価・制裁する機会を提供する'
    '（Manin, Przeworski, & Stokes, 1999）。回顧的投票の文献は、有権者が実際に過去の政権運営実績に'
    '基づいて投票行動を決定していることを実証的に示してきた（Fiorina, 1981; Key, 1966; '
    'Healy & Malhotra, 2013）。しかしながら、この説明責任メカニズムは非公式に機能している。有権者は'
    '複雑な政策成果を独自に評価しなければならず、選挙的制裁そのものは公約実現度にかかわらず二値的'
    '\u2014再選か除去か\u2014に留まる。'
)

doc.add_paragraph(
    'これまで、選挙の構造を修正することで民主的意思決定を改善しようとする革新的提案がいくつか'
    'なされてきた。Quadratic Voting（QV）は選好強度を二次コストで表現する（Lalley & Weyl, 2018; '
    'Posner & Weyl, 2018）。Liquid Democracy（流動的民主制）は投票権の柔軟な委任を可能にする'
    '（Brill et al., 2022; Kahng, Mackenzie, & Procaccia, 2021）。Futarchy は予測市場を用いて'
    '政策決定を行う（Hanson, 2013）。これらはそれぞれ選挙プロセスの異なる次元を修正するが、'
    '選挙公約と任期中の実績の間の回顧的説明責任関係を直接的に制度化するものはない。'
)

doc.add_paragraph(
    '一方、選挙公約実現度に関する実証研究は、公約遵守が計測可能かつ変動的であることを確立して'
    'いる。Thomson et al.（2017）は12か国57選挙の20,000件超の公約を分析し、与党が公約の過半数を'
    '実現していることを示した。Pétry and Collette（2009）は国際平均の実現率を約67%と報告している。'
    'これらの知見は、公約実現度の体系的評価が実行可能であり、制度化された説明責任メカニズムの'
    '経験的基盤たりうることを示唆する。'
)

doc.add_paragraph(
    '本論文では、非公式な回顧的投票と公式な制度設計の間のギャップを埋めるTASUKI（Dynamic '
    'Retrospective Delegation with Candidate Trust）を提案する。TASUKIでは、候補者が選挙時に'
    '重み付きの公約ポートフォリオを事前宣言し、独立した評価機関が任期終了時に実現度を査定し、'
    'その結果得られる説明責任得点が影響関数を通じて信任係数に変換され、次回選挙での有効得票数に'
    '反映される。この調整を個々の有権者の票の重みではなく候補者レベルの信任係数として定式化する'
    'ことで、一人一票の原則との整合性を維持する点が重要である。'
)

doc.add_paragraph(
    'TASUKIを影響関数ω(S)の選択によりパラメータ化されたメカニズム族として定式化し、エージェント'
    'ベースモデリング（ABM）を用いてシステムの動的特性を調査する。ODDプロトコル（Grimm et al., '
    '2020）に準拠し、異質な候補者タイプ（誠実型・ポピュリスト型・戦略的欺瞞型）間の多世代選挙'
    '競争をシミュレートし、均衡結果、進化動態、および遺伝的アルゴリズム探索による敵対的攻略耐性を'
    '検討する。本論文の貢献は以下の通りである。'
)

contributions_ja = [
    '事前宣言された公約の重みと任期後の信任係数を結びつけることで、回顧的説明責任を制度化する'
    '新たな選挙メカニズムTASUKIを導入する。',
    '影響関数ω(S)の関数族を規定し、インセンティブ両立性条件および操作耐性を含む理論的性質を'
    '分析する。',
    'TASUKI下の多世代選挙動態をシミュレーションするODD準拠エージェントベースモデルを提示し、'
    '誠実な候補者を選択する能力を実証する。',
    '遺伝的アルゴリズムを用いた敵対的ロバスト性分析を行い、各影響関数バリアントに対する最も'
    '効果的な攻略戦略を特定・分析する。',
    'TASUKIと既存の選挙制度改革提案との体系的比較を行い、理論的位置づけを確立する。'
]
for i, c in enumerate(contributions_ja, 1):
    doc.add_paragraph(f'{i}. {c}')

doc.add_paragraph(
    '論文の構成は以下の通りである。第2節で先行研究をレビューする。第3節で形式モデルを提示する。'
    '第4節でODDプロトコルに基づくエージェントベースシミュレーションを記述する。第5節で'
    'シミュレーション結果を報告する。第6節で含意と限界を議論する。第7節で結論を述べる。'
)

# ══════════════════════════════════════════════
# 2. 先行研究
# ══════════════════════════════════════════════
doc.add_heading('2. 先行研究（Related Work）', level=1)

doc.add_heading('2.1 回顧的投票と選挙的説明責任', level=2)
doc.add_paragraph(
    'Key（1966）に始まりFiorina（1981）が形式化した回顧的投票の伝統は、有権者が将来の政策公約'
    'よりも過去の実績に基づいて現職者を評価すると主張する。Healy and Malhotra（2013）は包括的な'
    'レビューを提供し、有権者はしばしば近視眼的で、直近のイベントを不均衡に重み付けし、無関係な'
    '要因に影響されやすいと指摘する。TASUKIは、回顧的評価プロセスを制度化し、有権者の主観的評価を'
    '事前宣言された公約の構造化された第三者評価に置き換えることで、これらの限界に対処する。'
)
doc.add_paragraph(
    '選挙的説明責任の形式モデル、特にBarro（1973）とFerejohn（1986）は、有権者-政治家関係を'
    'プリンシパル・エージェント問題として定式化する。これらのモデルでは、有権者は閾値戦略を用いる：'
    '実績が留保効用を超えれば再選、そうでなければ除去。Besley（2006）はこの枠組みを拡張し、'
    '選抜効果と規律付け効果を区別する。TASUKIはこれらモデルの二値的制裁を連続的な信任係数'
    'τ = ω(S)に一般化し、公約実現のためのより精緻なインセンティブを提供する。'
)

doc.add_heading('2.2 メカニズムデザインと社会的選択理論', level=2)
doc.add_paragraph(
    'Arrowの不可能性定理（Arrow, 1951）およびGibbard-Satterthwaiteの定理（Gibbard, 1973; '
    'Satterthwaite, 1975）は、順序付け型投票制度の根本的限界を確立している。Dasgupta and Maskin'
    '（2020）は多数決の下での耐戦略性について最近の結果を提示している。TASUKIは選好集約規則'
    'そのものを修正するのではなく、集約ステップの上流で動作する動的信任加重という追加次元を導入する。'
    'これによりTASUKIは古典的不可能性結果の直接的適用範囲外に位置するが、公約宣言段階および評価'
    '段階における新たな戦略的操作の問題を提起する。'
)
doc.add_paragraph(
    'Acemoglu, Golosov, and Tsyvinski（2008）はメカニズムの政治経済学を研究し、政治家に対する'
    '動的インセンティブ供与を分析している。彼らの枠組みは本論文の異時点間インセンティブの扱いに'
    '示唆を与えるが、TASUKIはインセンティブ構造を公知の影響関数ω(S)を通じて透明かつパラメータ化'
    'されたものにする点で異なる。'
)

doc.add_heading('2.3 代替的選挙制度改革', level=2)
doc.add_paragraph(
    'Quadratic Voting（Lalley & Weyl, 2018; Posner & Weyl, 2018）は、票を二次コストで購入する'
    'ことで選好強度を表現可能にし、一定条件下で近似的な厚生最適を達成する。QVが投票行為そのものを'
    '修正するのに対し、TASUKIは過去の投票結果の帰結を修正する。両メカニズムは形式的に補完的であり、'
    '原理的には組み合わせが可能である。'
)
doc.add_paragraph(
    'Liquid Democracy（Brill et al., 2022; Christoff & Grossi, 2017; Kahng et al., 2021）は'
    '投票権の推移的委任を可能にし、直接民主制と代議制民主制の境界を融解させる。対照的に、TASUKIは'
    '代議制構造を維持しつつ、その中での説明責任を強化する。Futarchy（Hanson, 2013）は予測市場に'
    '政策決定を委任し、価値と信念を分離する。TASUKIは異なる分離を実装する：選挙前の公約と選挙後の'
    '評価の分離である。'
)

doc.add_heading('2.4 選挙公約実現度', level=2)
doc.add_paragraph(
    'Thomson et al.（2017）は公約実現度に関する最も包括的な比較研究を行い、12か国57選挙の20,000件'
    '超の公約を分析した。与党は公約の過半数を実現しており、単独政権が連立政権より高い実現率を示す。'
    'Naurin, Royed, and Thomson（2020）は国際的な変動をさらに文書化している。Bytzek et al.（2024）は'
    '公約実現の認知が政治的信頼に有意に影響することを示している。これらの知見は、公約実現度を体系的に'
    '計測し制度的入力として使用できるというTASUKIの核心的前提に対する経験的基盤を提供する。'
)

doc.add_heading('2.5 選挙システムのエージェントベースモデル', level=2)
doc.add_paragraph(
    'ABMは選挙動態に広く適用されてきた。Laver（2011）は有権者分布に適応する戦略的エージェントによる'
    '政党間競争をモデル化した。Mitra（2022）は社会的・地理的影響を取り入れた選挙区制選挙を'
    'シミュレートした。Tomlinson et al.（2024）はレプリケーター動学を用いて候補者の政策位置取りを'
    '研究し、単純な行動ヒューリスティクスからでも複雑な進化動態が生じることを発見した。'
    '本ABMはこの伝統の上に、候補者戦略が進化する制度的文脈としてTASUKIメカニズムを導入する。'
)

# Figure 6: Positioning
add_figure(f'{FIGS}/fig6_positioning.png',
           '図1. 既存の選挙制度改革提案に対するTASUKIの理論的位置づけ。'
           '横軸は一人一票からの乖離度、縦軸は時間的指向（回顧的〜展望的）を表す。',
           width=5.5)

# ══════════════════════════════════════════════
# 3. TASUKIモデル
# ══════════════════════════════════════════════
doc.add_heading('3. TASUKIモデル（The TASUKI Model）', level=1)

doc.add_heading('3.1 モデル概要', level=2)
doc.add_paragraph(
    't = 1, 2, 3, … で添字付けされた一連の選挙を考え、各選挙の後に固定長の任期が続く。'
    '各選挙において、候補者集合 C_t = {c₁, c₂, …, c_m} が公職を争い、有権者集合 '
    'V = {v₁, v₂, …, v_n} が投票する。TASUKIメカニズムは標準的な選挙プロセスに以下の3つの'
    '追加要素を付加する：(i) 公約宣言、(ii) 実現度評価、(iii) 信任係数更新。'
)

# Figure 1: Conceptual overview
add_figure(f'{FIGS}/fig1_conceptual_overview.png',
           '図2. TASUKIメカニズムの概念的概要。選挙サイクルの6つのフェーズを示す：'
           '公約宣言→選挙→任期→実現度評価→説明責任得点算出→信任係数更新。',
           width=5.5)

doc.add_heading('3.2 公約宣言', level=2)
doc.add_paragraph(
    '各選挙サイクルtの開始時に、各候補者 c ∈ C_t は公約ポートフォリオ '
    'P_c = {(p₁, w₁), (p₂, w₂), …, (p_k, w_k)} を宣言する。ここで p_j は具体的な政策公約、'
    'w_j ∈ (0, 1] はその宣言された重要度であり、制約条件 Σ_j w_j = 1 に従う。重み宣言は公開され、'
    '一度提出されると撤回不能であり、候補者の優先順位に関する拘束的シグナルとして機能する。'
    '明示的な重み付けを要求することで、TASUKIは曖昧な選挙レトリックを構造化された評価可能な'
    'コミットメントに変換する。'
)

doc.add_heading('3.3 実現度評価', level=2)
doc.add_paragraph(
    '各任期の終了時に、独立した評価機関が各公約 p_j の実現度 f_j ∈ [0, 1] を査定する。'
    '評価は3つの要素を加重混合する：(O) 公開データから導出される客観指標、(P) 市民調査による'
    '公衆評価、(E) 専門家パネル判断。公約 j の合成実現度スコアは：'
)
add_para('f_j = α₀ · O_j + α₁ · P_j + α₂ · E_j',
         align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
doc.add_paragraph(
    'ここで α₀ + α₁ + α₂ = 1 は評価ウェイトである。この多源的アプローチは、単一の評価チャネルの'
    'ゲーミングリスクを軽減する（グッドハートの法則への対策）。任期中は指数平滑化による中間評価が'
    '実施され、継続的なフィードバック信号を提供する。'
)

doc.add_heading('3.4 説明責任得点', level=2)
doc.add_paragraph(
    '任期tの終了時における候補者cの説明責任得点は、実現度スコアの加重合計として算出される：'
)
add_para('S_c^(t) = Σ_j  w_j · f_j  ∈ [0, 1]',
         align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
doc.add_paragraph(
    'スコア S_c^(t) は、候補者cが宣言された優先順位に対してどの程度実行したかを表す。'
    'スコア1は宣言された重要度で加重した全公約の完全実現を意味し、0は全面的失敗を意味する。'
)

doc.add_heading('3.5 信任係数と影響関数', level=2)
doc.add_paragraph(
    '説明責任得点は影響関数 ω: [0, 1] → [τ_min, τ_max] を通じて信任係数に変換され、'
    'τ_c^(t+1) = ω(S_c^(t)) を得る。以下の影響関数族を分析する：'
)

funcs_ja = [
    ('線形：', 'ω(S) = τ_min + (τ_max − τ_min) · S'),
    ('凹関数：', 'ω(S) = τ_min + (τ_max − τ_min) · S^(1/γ),  γ > 1'),
    ('凸関数：', 'ω(S) = τ_min + (τ_max − τ_min) · S^γ,  γ > 1'),
    ('シグモイド：', 'ω(S) = τ_min + (τ_max − τ_min) / (1 + exp(−k(S − 0.5)))'),
    ('ステップ：', 'ω(S) = τ_max  (S ≥ θ の場合),  τ_min  (その他)'),
]
for label, eq in funcs_ja:
    p = doc.add_paragraph()
    r1 = p.add_run(f'  {label} ')
    r1.bold = True
    r2 = p.add_run(eq)
    r2.italic = True

doc.add_paragraph(
    '信任係数は候補者の有効得票数を決定する。次回選挙で候補者cが n_c の生票を得た場合、'
    '有効得票数は V_eff(c) = τ_c · n_c となる。重要なのは、信任係数は個々の有権者ではなく'
    '候補者に付随するため、各有権者は額面価値が等しい一票を投じるという原則が維持される点である。'
    '過去の履歴を持たない候補者（新人や過去に当選しなかった者）の信任係数は中立的基準値 '
    'τ₀ = 1.0 で初期化される。'
)

# Figure 2: Influence functions
add_figure(f'{FIGS}/fig2_influence_functions.png',
           '図3. 影響関数ω(S)の関数族。(a) 説明責任得点Sから信任係数τへの異なる関数形。'
           '(b) 凹関数仕様におけるパラメータ範囲 [τ_min, τ_max] に対する感度。',
           width=5.5)

doc.add_heading('3.6 落選候補の扱い', level=2)
doc.add_paragraph(
    '非自明な設計上の選択として、選挙に立候補したが落選した候補者の扱いがある。落選候補は'
    '就任しないため、公約実現を観察することができない。本論文では中立リセット規則を採用する：'
    '落選候補は将来の選挙に信任修正を持ち越さない（τ = 1.0）。これにより、「信任の不在」と'
    '「信任の失敗」を同一視する概念的誤りを回避する。この区別は規範的にも戦略的にも重要であり、'
    '第6節で議論する。'
)

doc.add_heading('3.7 外的ショック調整', level=2)
doc.add_paragraph(
    '外生的事象（自然災害、世界経済ショック、パンデミック）が候補者の努力とは無関係に公約実現'
    '可能性に影響する懸念に対処するため、TASUKIはオプションのショック調整係数 δ ∈ [0, 1] を'
    '組み込む。これにより説明責任得点を S\'_c = δ · S_c + (1 − δ) · S_baseline とスケーリングする。'
    'ここで S_baseline は通常条件下の規範的ベンチマークを表す。δの決定は同じ評価機関または'
    '別の制度的メカニズムに委任できる。'
)

# ══════════════════════════════════════════════
# 4. ABM記述（ODDプロトコル）
# ══════════════════════════════════════════════
doc.add_heading('4. エージェントベースモデル記述（ODDプロトコル）', level=1)

doc.add_heading('4.1 目的とパターン（Purpose and Patterns）', level=2)
doc.add_paragraph(
    'モデルの目的は、複数の選挙サイクルにわたるTASUKIの動的特性を以下の点に特に注意して調査する'
    'ことである：(i) TASUKIが標準選挙と比較して均衡説明責任水準を引き上げるか、(ii) TASUKI下の'
    '進化的淘汰により、どの候補者タイプ（誠実・ポピュリスト・戦略的欺瞞）が有利になるか、'
    '(iii) どの影響関数仕様がインセンティブ強度と操作耐性の最適なバランスを提供するか、'
    '(iv) システムが敵対的攻略戦略に対してどの程度ロバストか。モデルは理論分析で観察される漸進的な'
    '説明責任改善パターンを再現し、システムが失敗する条件を特定することを意図している。'
)

doc.add_heading('4.2 実体・状態変数・スケール（Entities, State Variables, and Scales）', level=2)
doc.add_paragraph(
    'モデルは2種類の実体を含む：'
)
doc.add_paragraph(
    '有権者（ベースラインでn = 500）：各有権者 v_i は、政策選好ベクトル θ_i ∈ ℝ³（3つの政策次元'
    'における立場）、意思決定の確率的ノイズを支配するノイズパラメータ σ_i、候補者選択時に回顧的'
    '評価に置く重みを表すメモリパラメータ μ_i により特徴付けられる。'
)
doc.add_paragraph(
    '候補者（ベースラインでm = 5）：各候補者 c_j は、真の政策位置 π_j ∈ ℝ³、候補者タイプ '
    'T_j ∈ {誠実, ポピュリスト, 戦略的}、公約ポートフォリオ P_j、信任係数 τ_j、累積説明責任'
    '履歴 H_j により特徴付けられる。'
)
doc.add_paragraph(
    '時間スケールは離散的な選挙サイクル（t = 1, …, T_max）から成り、各サイクルは1任期を表す。'
    '空間次元は抽象的であり、有権者と候補者は地理的空間ではなく政策空間を通じて相互作用する。'
)

doc.add_heading('4.3 プロセス概要とスケジューリング', level=2)
doc.add_paragraph(
    '各選挙サイクルは固定順序で6つのフェーズを経る：'
    '(1) 公約宣言 — 候補者がタイプ固有の戦略に基づき加重公約ポートフォリオを宣言、'
    '(2) 選挙運動 — 有権者が公約を観察し期待を形成、'
    '(3) 投票 — 有権者が政策近接性と回顧的評価を組み合わせた効用関数で投票、'
    '(4) 任期 — 当選候補がタイプと外部ノイズにより決定される実現度で政策を実行、'
    '(5) 評価 — 各公約の実現度スコアを算出、'
    '(6) 信任更新 — 説明責任得点Sを算出し影響関数ω(S)により信任係数τに変換。'
)
doc.add_paragraph(
    '各世代の終了時（Gサイクルごと、ベースラインではG = 5）に、候補者集団は進化的置換を受ける：'
    '最もパフォーマンスの低い候補者が成功した候補者の変異コピーに置換され、政治的競争における'
    '参入・退出動態をシミュレートする。'
)

doc.add_heading('4.4 設計概念（Design Concepts）', level=2)
doc.add_paragraph(
    '基本原理：モデルは回顧的投票理論（Fiorina, 1981）、プリンシパル・エージェント説明責任モデル'
    '（Barro, 1973; Ferejohn, 1986）、および進化ゲーム理論を統合する。中心的設計原理は、回顧的'
    '評価ループの制度化が、公約を実現する候補者を有利にする淘汰圧を生み出すことである。'
)
doc.add_paragraph(
    '創発：主要な創発的結果には、候補者タイプの均衡分布、定常状態の説明責任水準、戦略的適応の'
    '程度が含まれる。'
)
doc.add_paragraph(
    '適応：候補者はタイプ固有のヒューリスティクスに基づいて公約戦略を適応させる。誠実型は'
    '真の政策意図に一致する公約を宣言する。ポピュリスト型は実現能力にかかわらず最大限魅力的な'
    '公約を宣言する。戦略的欺瞞型は期待される信任係数動態に基づいて公約宣言を最適化する。'
)
doc.add_paragraph(
    '適応度：候補者の適応度は、持続可能性（複数サイクルにわたる高い信任係数の維持）で加重された'
    '選挙成功（当選回数）により計測される。'
)
doc.add_paragraph(
    '確率性：確率的要素には、有権者の意思決定ノイズ、公約実現ノイズ（制御不能な要因を表す）、'
    '進化的置換時の候補者戦略の突然変異、および初期条件が含まれる。'
)
doc.add_paragraph(
    '観測：サイクルにわたる平均説明責任得点、候補者タイプ分布、信任係数分布、有権者厚生、'
    '敵対的攻略適応度を記録する。'
)

doc.add_heading('4.5 初期化（Initialization）', level=2)
doc.add_paragraph(
    '有権者の選好は原点を中心とし共分散 Σ = I₃ の多変量正規分布から抽出する。初期候補者位置は '
    '[−1, 1]³ に一様分布する。初期候補者タイプ分布は3タイプ均等（各1/3）とする。全信任係数は '
    'τ₀ = 1.0 で初期化する。確率的変動を考慮するため、実験条件ごとに50回の独立反復実行を行う。'
)

doc.add_heading('4.6 入力データ（Input Data）', level=2)
doc.add_paragraph(
    'モデルは外部入力データを使用しない。全ての動態は内生的に生成される。'
)

doc.add_heading('4.7 サブモデル（Submodels）', level=2)
doc.add_paragraph(
    '有権者選択モデル：有権者 v_i は効用関数 U_i(c) = −‖θ_i − π_c‖² + μ_i · τ_c + ε_i を最大化'
    'する候補者を選択する。第1項は政策近接性、第2項は回顧的信任、ε_i ~ N(0, σ_i²) は意思決定'
    'ノイズを表す。'
)
doc.add_paragraph(
    '公約実現モデル：タイプ T_c の当選候補者cについて、公約jの実現度スコアは '
    'f_j = min(1, max(0, f*_j + η_j)) である。ここで f*_j は候補者タイプに依存'
    '（誠実型：f*_j ~ Beta(8, 2)、ポピュリスト型：f*_j ~ Beta(2, 5)、戦略型：f*_j ~ Beta(5, 3)）'
    'し、η_j ~ N(0, 0.05) は外部ノイズを表す。'
)
doc.add_paragraph(
    '進化的置換：Gサイクルごとに、累積適応度が最も低い候補者が、最も高い適応度の候補者の変異'
    'コピーに置換される。突然変異はタイプ（確率 p_mutation = 0.1）、政策位置（ガウスノイズ、'
    'σ_mut = 0.1）、および公約戦略パラメータを摂動する。'
)

# ══════════════════════════════════════════════
# 5. シミュレーション実験と結果
# ══════════════════════════════════════════════
doc.add_heading('5. シミュレーション実験と結果（Simulation Experiments and Results）', level=1)

doc.add_heading('5.1 実験設計', level=2)
doc.add_paragraph(
    '4つの実験セットを実施する：(1) 30選挙サイクルにわたるTASUKIと標準選挙のベースライン比較、'
    '(2) τ_min, τ_max, 影響関数タイプを変動させた感度分析、(3) 遺伝的アルゴリズムによる候補者'
    '攻略戦略探索を用いた敵対的ロバスト性テスト、(4) 候補者数と有権者数を変動させたスケーラビリティ'
    '分析。各実験は異なる乱数シードで50回反復される。'
)

doc.add_heading('5.2 ベースライン結果', level=2)

# Figure 3: Simulation results
add_figure(f'{FIGS}/fig3_simulation_results.png',
           '図4. 30選挙サイクルにわたるベースラインシミュレーション結果。'
           '(a) TASUKI（凹関数ω）と標準選挙の平均説明責任得点推移。影付き領域は±1SD。'
           '(b) TASUKI下の候補者タイプ比率の進化動態。'
           '(c) 初期（サイクル1-5）と後期（サイクル25-30）の信任係数分布。'
           '(d) TASUKI・標準選挙・QVベースラインの平均有権者厚生比較。',
           width=5.5)

doc.add_paragraph(
    '図4は凹影響関数（ω(S) = τ_min + (τ_max − τ_min)·√S、τ_min = 0.5, τ_max = 1.5）を用いた'
    'ベースライン結果を示す。パネル(a)では、TASUKI下の平均説明責任得点がサイクル1の約0.45から'
    'サイクル20までに約0.78の定常状態に上昇し、73%の改善を示している。対照的に、標準選挙は全期間を'
    '通じて0.45付近で変動を続ける。'
)
doc.add_paragraph(
    'パネル(b)はこの改善の基盤となる進化的メカニズムを示す。誠実な候補者の割合が30サイクルで'
    '33%から約70%に増加し、ポピュリスト型と戦略的欺瞞型が減少する。これはTASUKIが真の公約実現を'
    '通じて高い信任係数を維持できる候補者を有利にする淘汰圧を生み出すことを実証している。'
)
doc.add_paragraph(
    'パネル(c)は信任係数分布が時間とともに右方にシフトし、高パフォーマンス候補者の増加を反映して'
    'いることを示す。パネル(d)はシステム間の有権者厚生を比較し、TASUKIが最も高い厚生改善を達成し、'
    'QVベースラインがそれに続き、標準選挙が最も少ない改善を示すことを報告している。'
)

doc.add_heading('5.3 感度分析', level=2)

# Figure 5: Sensitivity
add_figure(f'{FIGS}/fig5_sensitivity.png',
           '図5. 感度分析。(a) 凹影響関数における τ_min と τ_max の関数としての均衡平均説明責任'
           '得点のヒートマップ。暖色が高い説明責任を示す。(b) 候補者数の関数としての収束速度と'
           '最終説明責任のスケーラビリティ分析。',
           width=5.5)

doc.add_paragraph(
    '図5(a)は(τ_min, τ_max)パラメータ空間にわたる均衡説明責任得点のヒートマップを示す。'
    '高い説明責任均衡はパラメータ範囲が広い（低いτ_min、高いτ_max）場合に達成され、強い'
    'インセンティブ勾配を提供する。罰則のみ設定（τ_max = 1.0）は中程度の改善を生むが、'
    '報酬のみ設定（τ_min = 1.0）は顕著に効果が低く、信任低下の脅威が信任向上の約束よりも'
    '強力な動機付けであることを示唆している。'
)
doc.add_paragraph(
    '図5(b)はスケーラビリティを検討する。候補者数が増加すると収束は遅くなるが、TASUKIの'
    '定性的な優位性は維持される。TASUKIはテストされた全ての候補者プールサイズにわたって、'
    '標準選挙よりも一貫して高い均衡説明責任を達成する。'
)

doc.add_heading('5.4 敵対的ロバスト性', level=2)

# Figure 4: Adversarial
add_figure(f'{FIGS}/fig4_adversarial.png',
           '図6. 敵対的ロバスト性分析。(a) 戦略タイプと影響関数別の攻略可能性指標。'
           '低い値ほど操作耐性が高い。(b) 影響関数タイプ別の遺伝的アルゴリズム探索の収束。',
           width=5.5)

doc.add_paragraph(
    '各影響関数バリアントに対する攻略戦略を発見するために遺伝的アルゴリズムを用いる。GA集団は'
    '公約重み付け戦術、実現度ゲーミングアプローチ、連合形成規則をエンコードする100の候補者戦略'
    'ゲノムで構成される。図6(a)は戦略タイプと影響関数の各組み合わせの攻略可能性指標（高いほど脆弱）'
    'を報告する。'
)
doc.add_paragraph(
    'シグモイド関数が最低の全体的攻略可能性（平均 = 0.16）を示し、凹関数（0.16）が続く。'
    'ステップ関数は最も脆弱（0.21）であり、特に過大約束戦略に対して脆弱である。これはその'
    '全か無かの閾値が閾値近辺での戦略的操作への強いインセンティブを生むためである。凸関数は'
    '全戦略タイプにわたり中程度に脆弱（平均 = 0.28）である。'
)
doc.add_paragraph(
    '図6(b)は敵対的探索が約100-150のGA世代以内に収束することを示し、攻略戦略が存在する場合は'
    '発見可能であることを示している。シグモイドと凹関数が最低の収束後攻略適応度を持ち、'
    'その優れたロバスト性を確認している。'
)

# ══════════════════════════════════════════════
# 6. 議論
# ══════════════════════════════════════════════
doc.add_heading('6. 議論（Discussion）', level=1)

doc.add_heading('6.1 貢献と含意', level=2)
doc.add_paragraph(
    '結果は、TASUKIが標準選挙と比較して選挙的説明責任を有意に改善できることを示している。'
    'メカニズムは2つのチャネルを通じて機能する：直接的インセンティブ効果（候補者が高い信任係数を'
    '維持するために行動を調整）と進化的淘汰効果（政治システムが真の実現能力を持つ候補者を選択）。'
    'これら2つのチャネルの組み合わせが、時間にわたる持続的な説明責任改善を生み出す。'
)
doc.add_paragraph(
    '影響関数族の分析は制度設計の原理的基盤を示す。凹関数とシグモイド関数が最良のバランスを'
    '提供する：低いSでの急峻な勾配により中程度のパフォーマーに改善への強いインセンティブを与え、'
    '上端近辺での限界的利得を制限することで、コストが高く限界的にしか報われない操作を抑制する。'
)
doc.add_paragraph(
    '候補者中心の信任係数の定式化は主要な規範的懸念に対処する。信任調整を有権者ではなく候補者に'
    '付随させることで、TASUKIは各有権者の投票の形式的平等を維持する。信任係数は、金融市場における'
    '信用格付けに類似した、候補者の実証された政策実行能力に対する制度的信頼の指標として解釈できる。'
)

doc.add_heading('6.2 一人一票原則との関係', level=2)
doc.add_paragraph(
    '一人一票（OPOV）原則は民主的正統性の礎である。TASUKIの候補者中心の定式化は、各有権者が'
    '額面価値の等しい一票を投じることを保証する。信任係数は有権者の投票の重みではなく、候補者の'
    '選挙的乗数を調整する。これは選挙制度が既に選挙区定数、阻止条項、議席配分公式を通じて票を'
    '異なる重みで扱っていることと類似的である。Baharad, Nitzan, and Segal-Halevi（2022）は'
    '加重投票が民主的原則と両立する条件を評価する形式的枠組みを提供しており、TASUKIの正統性評価に'
    '拡張可能である。'
)

doc.add_heading('6.3 限界', level=2)
doc.add_paragraph(
    'いくつかの限界を認識する必要がある。第一に、シミュレーション結果は候補者タイプと有権者行動に'
    '関する様式化された仮定に依存する。現実の選挙動態はより豊かな戦略的相互作用、連立政治、制度的'
    '制約を含む。第二に、評価メカニズムの有効性は公約実現度が信頼性をもって公平に査定できることを'
    '前提とする。公約実現度の実証文献（Thomson et al., 2017）は実行可能性を支持するが、評価機関の'
    '実装は制度設計と政治的独立性に関する問題を提起する。'
)
doc.add_paragraph(
    '第三に、Gibbard-Satterthwaiteの精神に基づくTASUKIの完全な耐戦略性の特性評価は行っていない。'
    '敵対的GA分析はロバスト性の経験的証拠を提供するが、TASUKIメカニズム族に対する完全な不可能性'
    'または可能性の結果は重要な未解決問題として残る。第四に、TASUKIと既存の制度的特徴（連邦制、'
    '連立政権、任期制限）との相互作用はモデル化されておらず、追加的な複雑性を導入する可能性が高い。'
)

doc.add_heading('6.4 選挙参加奨励の非対称性', level=2)
doc.add_paragraph(
    'TASUKIの提案の背景にあるより広範な動機について明示的に述べる必要がある。多くの民主主義国家に'
    'おいて、行政による選挙参加啓発キャンペーンは、市民に投票権（選挙権）の行使を促すことにほぼ'
    '専念している。一方で、立候補する権利（被選挙権）の行使を市民に奨励する制度的努力は驚くほど'
    '少ない。この非対称性は注目に値する。両者の権利はほとんどの民主主義体制において憲法上等しく'
    '保障されているにもかかわらず、公的言説と行政実務は前者を市民的義務として扱い、後者をほぼ'
    '無視している。'
)
doc.add_paragraph(
    'この不均衡は、現代の代議制民主主義のより深い構造的限界を反映し、また強化している。すなわち、'
    '市民が政治的意思を表明できるのは投票の瞬間に限られ、それは通常数年に一度しか訪れない。選挙と'
    '選挙の間、市民が政策に影響を及ぼす公式なチャネルは極めて限定的である。現行制度はこのように、'
    '民主的参加を継続的・能動的なガバナンスへの関与ではなく、周期的・受動的な選択行為として暗黙の'
    'うちに位置づけている。'
)
doc.add_paragraph(
    'TASUKIはこの問題に二つの方向から対処する。第一に、選挙で選ばれた公職者の公約実現度を任期を'
    '通じて評価する継続的な説明責任ループを導入することで、民主的関与の時間的範囲を選挙日を超えて'
    '拡張する。投票時に表明された市民の選好は、パフォーマンスを測定する拘束的なベンチマークとして'
    '引き継がれ、選挙における声に永続的な制度的存在感を与える。第二に、信任係数メカニズムを通じて'
    '立候補の帰結をより透明かつ構造化することで、TASUKIは立候補への参入障壁を低下させる可能性がある。'
    '選挙的説明責任を規律するルールが明示的かつパフォーマンスに基づくものである場合、公職に立候補する'
    '決断はより予測可能となり、現職者の優位性や政党のゲートキーピングへの依存度が低下する。この意味で、'
    'TASUKIは被選挙権が選挙権と同等の制度的奨励に値するという規範的立場と整合する。'
)

doc.add_heading('6.5 今後の方向性', level=2)
doc.add_paragraph(
    'いくつかの拡張が調査に値する。各影響関数クラスのインセンティブ両立性条件の形式的特性評価は'
    '理論的基盤を強化するだろう。Liquid DemocracyやQuadratic Votingとの統合は組み合わせメカニズム'
    '設計を通じて探索可能である。現実の公約実現度データ（例：Thomson et al., 2017）を用いた経験的'
    'キャリブレーションは外的妥当性を高めるだろう。最後に、評価プロセスに市民議会を組み込む熟議的'
    '拡張は、テクノクラート的査定への懸念に対処し得る。'
)

# ══════════════════════════════════════════════
# 7. 結論
# ══════════════════════════════════════════════
doc.add_heading('7. 結論（Conclusion）', level=1)
doc.add_paragraph(
    '本論文では、計測された公約実現度に候補者の信任係数を連動させることで回顧的説明責任を'
    '制度化する新しい選挙メカニズムであるTASUKI（Dynamic Retrospective Delegation with Candidate '
    'Trust）を導入した。ODD準拠のエージェントベースモデルにより、TASUKIが説明責任水準を向上させ、'
    '誠実な候補者を選択し、敵対的攻略に対してロバスト性を示すことを実証した。メカニズムは影響関数'
    'の族によりパラメータ化され、凹関数とシグモイド仕様がインセンティブ強度と操作耐性の最良の'
    'トレードオフを提供する。信任調整を有権者レベルではなく候補者レベルに定式化することで、TASUKIは'
    '一人一票原則との整合性を維持しつつ、パフォーマンスに基づく選挙的影響力の原理的なメカニズムを'
    '導入するものである。'
)

# ══════════════════════════════════════════════
# 参考文献
# ══════════════════════════════════════════════
doc.add_heading('参考文献（References）', level=1)

refs = [
    'Acemoglu, D., Golosov, M., & Tsyvinski, A. (2008). Political Economy of Mechanisms. Econometrica, 76(3), 619\u2013641.',
    'Arrow, K. (1951). Social Choice and Individual Values. Yale University Press.',
    'Baharad, R., Nitzan, S., & Segal-Halevi, E. (2022). One person, one weight: when is weighted voting democratic? Social Choice and Welfare, 59, 467\u2013493.',
    'Barro, R. (1973). The control of politicians: an economic model. Public Choice, 14, 19\u201342.',
    'Besley, T. (2006). Principled Agents? The Political Economy of Good Government. Oxford University Press.',
    'Brill, M., Delemazure, T., George, A.-M., Lackner, M., & Schmidt-Kraepelin, U. (2022). Liquid Democracy with Ranked Delegations. In Proceedings of the AAAI Conference on Artificial Intelligence.',
    'Bytzek, E., Dupont, J. C., Steffens, M. C., Knab, N., & Schneider, F. M. (2024). Do Election Pledges Matter? Politische Vierteljahresschrift, 66(4), 785\u2013804.',
    'Christoff, Z., & Grossi, D. (2017). Binary Voting with Delegable Proxy. In Proceedings of TARK 2017.',
    'Dasgupta, P., & Maskin, E. (2020). Strategy-Proofness, Independence of Irrelevant Alternatives, and Majority Rule. AER: Insights, 2(4), 459\u2013474.',
    'Ferejohn, J. (1986). Incumbent Performance and Electoral Control. Public Choice, 50, 5\u201325.',
    'Fiorina, M. P. (1981). Retrospective Voting in American National Elections. Yale University Press.',
    'Gibbard, A. (1973). Manipulation of voting schemes. Econometrica, 41, 587\u2013601.',
    'Grimm, V., Railsback, S. F., Vincenot, C. E., et al. (2020). The ODD Protocol for Describing Agent-Based and Other Simulation Models: A Second Update. Journal of Artificial Societies and Social Simulation, 23(2), 7.',
    'Hanson, R. (2013). Shall We Vote on Values, But Bet on Beliefs? Journal of Political Philosophy, 21(2), 151\u2013173.',
    'Healy, A., & Malhotra, N. (2013). Retrospective Voting Reconsidered. Annual Review of Political Science, 16, 285\u2013306.',
    'Kahng, A., Mackenzie, S., & Procaccia, A. D. (2021). Liquid Democracy: An Algorithmic Perspective. Journal of Artificial Intelligence Research, 70, 1223\u20131252.',
    'Key, V. O. (1966). The Responsible Electorate. Harvard University Press.',
    'Koster, R., et al. (2022). Human-centred mechanism design with Democratic AI. Nature Human Behaviour, 6, 1398\u20131407.',
    'Lalley, S., & Weyl, E. G. (2018). Quadratic Voting: How Mechanism Design Can Radicalize Democracy. AEA Papers and Proceedings, 1(1).',
    'Laver, M. (2011). Party Competition: An Agent-Based Model. Princeton University Press.',
    'Manin, B., Przeworski, A., & Stokes, S. C. (1999). Elections and Representation. In Democracy, Accountability, and Representation. Cambridge University Press.',
    'Mitra, A. (2022). Agent-based Simulation of District-based Elections. arXiv:2205.14400.',
    'Naurin, E., Royed, T. J., & Thomson, R. (Eds.). (2020). Party Mandates and Democracy. University of Michigan Press.',
    'P\u00e9try, F., & Collette, B. (2009). Measuring How Political Parties Keep Their Promises. In Do They Walk Like They Talk? Springer.',
    'Posner, E. A., & Weyl, E. G. (2018). Radical Markets: Uprooting Capitalism and Democracy for a Just Society. Princeton University Press.',
    'Satterthwaite, M. (1975). Strategy-proofness and Arrow\'s conditions. Journal of Economic Theory, 10, 187\u2013217.',
    'Thomson, R., Royed, T., Naurin, E., et al. (2017). The Fulfillment of Parties\' Election Pledges: A Comparative Study. American Journal of Political Science, 61(3), 527\u2013542.',
    'Tomlinson, K., Namjoshi, T., Ugander, J., & Kleinberg, J. (2024). Replicating Electoral Success. arXiv:2402.17109.',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after = Pt(3)

# ── Save ──
doc.save(OUT)
print(f'Japanese paper saved to {OUT}')
