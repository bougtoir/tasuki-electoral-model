#!/usr/bin/env python3
"""Generate English JASSS paper as .docx with embedded color figures."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

FIGS = '/home/ubuntu/figures'
OUT = '/home/ubuntu/TASUKI_JASSS_English.docx'

doc = Document()

# ── Style Setup ──────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Arial'
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    if level == 1:
        hs.font.size = Pt(16)
    elif level == 2:
        hs.font.size = Pt(13)
    else:
        hs.font.size = Pt(11)

def add_para(text, bold=False, italic=False, align=None, size=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_figure(path, caption, width=5.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(10)
    cap.paragraph_format.space_after = Pt(12)

# ══════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════
add_para('', space_after=24)
add_para(
    'Trust-Adjusted Scoring with Unified Knowledge Integration (TASUKI):\n'
    'An Agent-Based Model of Accountability-Driven Electoral Reform',
    bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18, space_after=12
)
add_para('[Author Name Redacted for Peer Review]',
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=6)
add_para('[Institutional Affiliation Redacted]',
         italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, space_after=24)

# ── Abstract ──
doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    'Democratic elections serve as the primary mechanism through which citizens hold representatives '
    'accountable, yet the binary nature of electoral sanctions (re-election versus removal) provides '
    'only a coarse instrument for incentivizing policy fulfillment. We propose Dynamic Retrospective '
    'Delegation with Candidate Trust (TASUKI), a novel electoral mechanism in which candidates '
    'pre-declare weighted policy pledges, third-party evaluators assess fulfillment at term end, '
    'and the resulting accountability score modulates a candidate-level trust coefficient that '
    'influences their effective electoral support in subsequent elections. Using an agent-based model '
    'compliant with the ODD protocol, we simulate multi-generational electoral dynamics under TASUKI '
    'across a family of influence functions. Our results show that TASUKI (i) significantly raises '
    'mean accountability scores relative to standard elections, (ii) induces evolutionary selection '
    'pressure favoring sincere candidates over populist and deceptive strategists, and (iii) exhibits '
    'robustness to a range of adversarial exploitation strategies discovered via genetic algorithm '
    'search. Sensitivity analysis reveals that concave and sigmoid influence functions offer the best '
    'trade-off between incentive strength and resistance to manipulation. We discuss implications for '
    'institutional design and the compatibility of TASUKI with the one-person-one-vote principle '
    'through candidate-centric trust reinterpretation.'
)

add_para('Keywords: electoral accountability, agent-based model, mechanism design, '
         'retrospective voting, weighted trust, policy fulfillment, ODD protocol',
         italic=True, size=10, space_after=18)

# ══════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════
doc.add_heading('1. Introduction', level=1)

doc.add_paragraph(
    'Elections constitute the cornerstone of representative democracy, providing citizens with '
    'periodic opportunities to evaluate and sanction their representatives (Manin, Przeworski, & '
    'Stokes, 1999). The retrospective voting literature has documented that voters do, in practice, '
    'condition their electoral choices on perceived government performance (Fiorina, 1981; Key, 1966; '
    'Healy & Malhotra, 2013). However, this accountability mechanism operates informally: voters must '
    'independently assess complex policy outcomes, and the electoral sanction itself is binary\u2014'
    're-election or removal\u2014regardless of the degree of policy fulfillment or failure.'
)

doc.add_paragraph(
    'Several innovative proposals have sought to improve democratic decision-making by modifying '
    'the structure of elections. Quadratic Voting (QV) addresses the failure to capture preference '
    'intensity by allowing voters to purchase additional votes at quadratic cost (Lalley & Weyl, 2018; '
    'Posner & Weyl, 2018). Liquid democracy enables flexible delegation of voting rights (Brill et al., '
    '2022; Kahng, Mackenzie, & Procaccia, 2021). Futarchy separates value judgments from empirical '
    'beliefs by employing prediction markets for policy selection (Hanson, 2013). Each of these '
    'approaches modifies a different dimension of the electoral process, yet none directly '
    'institutionalizes the retrospective accountability relationship between campaign promises and '
    'post-election performance.'
)

doc.add_paragraph(
    'Meanwhile, the empirical literature on campaign pledge fulfillment has established that promise '
    'keeping is both measurable and variable. Thomson et al. (2017), analyzing over 20,000 pledges '
    'across 12 countries, found that governing parties fulfill a majority of their promises, with '
    'rates varying significantly based on institutional context. P\u00e9try and Collette (2009) report '
    'a cross-national average fulfillment rate of approximately 67%. These findings suggest that '
    'systematic evaluation of pledge fulfillment is feasible and could serve as the empirical '
    'foundation for an institutionalized accountability mechanism.'
)

doc.add_paragraph(
    'In this paper, we propose Dynamic Retrospective Delegation with Candidate Trust (TASUKI), '
    'an electoral mechanism that bridges the gap between informal retrospective voting and formal '
    'institutional design. Under TASUKI, candidates pre-declare a weighted portfolio of policy pledges '
    'at election time, an independent evaluation body assesses fulfillment at the end of the term, '
    'and the resulting accountability score is transformed via an influence function into a trust '
    'coefficient that modulates the candidate\'s effective electoral support in subsequent elections. '
    'Crucially, we frame this adjustment as a candidate-level trust coefficient rather than a '
    'modification of individual voter weights, thereby maintaining compatibility with the '
    'one-person-one-vote principle.'
)

doc.add_paragraph(
    'We formalize TASUKI as a family of mechanisms parameterized by the choice of influence function '
    '\u03c9(S), and employ agent-based modeling (ABM) to investigate the system\'s dynamic properties. '
    'Following the ODD protocol (Grimm et al., 2020), we simulate multi-generational electoral '
    'competition among heterogeneous candidate types (sincere, populist, and strategic-deceptive) and '
    'examine equilibrium outcomes, evolutionary dynamics, and robustness to adversarial exploitation '
    'via genetic algorithm search. Our contributions are as follows:'
)

# Numbered contributions
contributions = [
    'We introduce TASUKI, a formal electoral mechanism that institutionalizes retrospective '
    'accountability by linking pre-declared pledge weights to post-term trust coefficients.',
    'We specify a family of influence functions \u03c9(S) and analyze their theoretical properties, '
    'including incentive compatibility conditions and manipulation resistance.',
    'We present an ODD-compliant agent-based model simulating multi-generational electoral '
    'dynamics under TASUKI, demonstrating its capacity to select for sincere candidates.',
    'We conduct adversarial robustness analysis using genetic algorithms to identify and '
    'characterize the most effective exploitation strategies for each influence function variant.',
    'We provide a systematic comparison of TASUKI with related electoral reform proposals, '
    'establishing its unique position in the theoretical landscape.'
]
for i, c in enumerate(contributions, 1):
    doc.add_paragraph(f'{i}. {c}')

doc.add_paragraph(
    'The remainder of this paper is organized as follows. Section 2 reviews related work. '
    'Section 3 presents the formal model. Section 4 describes the agent-based simulation following '
    'the ODD protocol. Section 5 reports simulation results. Section 6 discusses implications and '
    'limitations. Section 7 concludes.'
)

# ══════════════════════════════════════════════
# 2. RELATED WORK
# ══════════════════════════════════════════════
doc.add_heading('2. Related Work', level=1)

doc.add_heading('2.1 Retrospective Voting and Electoral Accountability', level=2)
doc.add_paragraph(
    'The retrospective voting tradition, initiated by Key (1966) and formalized by Fiorina (1981), '
    'posits that voters evaluate incumbents based on past performance rather than prospective policy '
    'promises. Healy and Malhotra (2013) provide a comprehensive review, noting that voters are often '
    'myopic, weighting recent events disproportionately, and susceptible to irrelevant factors. '
    'TASUKI addresses these limitations by institutionalizing the retrospective evaluation process, '
    'replacing subjective voter assessments with structured third-party evaluation of pre-declared '
    'pledges.'
)
doc.add_paragraph(
    'Formal models of electoral accountability, notably Barro (1973) and Ferejohn (1986), frame the '
    'voter-politician relationship as a principal-agent problem. In these models, voters employ '
    'threshold strategies: re-elect the incumbent if performance exceeds a reservation utility, '
    'otherwise remove. Besley (2006) extends this framework to distinguish selection and disciplining '
    'effects. TASUKI generalizes the binary sanction of these models into a continuous trust '
    'coefficient \u03c4 = \u03c9(S), providing finer-grained incentives for policy fulfillment.'
)

doc.add_heading('2.2 Mechanism Design and Social Choice Theory', level=2)
doc.add_paragraph(
    'Arrow\'s impossibility theorem (Arrow, 1951) and the Gibbard-Satterthwaite theorem (Gibbard, '
    '1973; Satterthwaite, 1975) establish fundamental limits on ranked-choice voting systems. '
    'Dasgupta and Maskin (2020) provide recent results on strategy-proofness under majority rule. '
    'TASUKI does not modify the preference aggregation rule itself but introduces an additional '
    'dimension\u2014dynamic trust weighting\u2014that operates upstream of the aggregation step. This '
    'situates TASUKI outside the direct scope of classical impossibility results, though it raises '
    'new questions about strategic manipulation in the pledge-declaration and evaluation stages.'
)
doc.add_paragraph(
    'Acemoglu, Golosov, and Tsyvinski (2008) study the political economy of mechanisms, analyzing '
    'dynamic incentive provision for politicians. Their framework informs our treatment of '
    'intertemporal incentives, though TASUKI differs in making the incentive structure transparent '
    'and parameterized through the publicly known influence function \u03c9(S).'
)

doc.add_heading('2.3 Alternative Electoral Reforms', level=2)
doc.add_paragraph(
    'Quadratic Voting (Lalley & Weyl, 2018; Posner & Weyl, 2018) allows voters to express '
    'preference intensity by purchasing votes at quadratic cost, achieving approximate welfare '
    'optimality under certain conditions. While QV modifies the voting act itself, TASUKI modifies '
    'the consequences of past voting outcomes. The two mechanisms are formally complementary and '
    'could, in principle, be combined.'
)
doc.add_paragraph(
    'Liquid democracy (Brill et al., 2022; Christoff & Grossi, 2017; Kahng et al., 2021) enables '
    'transitive delegation of voting rights, dissolving the boundary between direct and '
    'representative democracy. In contrast, TASUKI preserves the representative structure while '
    'strengthening accountability within it. Futarchy (Hanson, 2013) delegates policy decisions to '
    'prediction markets, separating values from beliefs. TASUKI implements a different separation: '
    'pre-election pledges from post-election evaluation.'
)

doc.add_heading('2.4 Campaign Pledge Fulfillment', level=2)
doc.add_paragraph(
    'Thomson et al. (2017) conduct the most comprehensive comparative study of pledge fulfillment, '
    'analyzing over 20,000 pledges across 57 election campaigns in 12 countries. They find that '
    'governing parties fulfill a majority of pledges, with single-party governments achieving higher '
    'rates than coalitions. Naurin, Royed, and Thomson (2020) further document cross-national '
    'variation. Bytzek et al. (2024) demonstrate that pledge fulfillment perceptions significantly '
    'affect political trust. These findings provide the empirical foundation for TASUKI\'s core '
    'assumption that pledge fulfillment can be systematically measured and used as an institutional '
    'input.'
)

doc.add_heading('2.5 Agent-Based Models of Electoral Systems', level=2)
doc.add_paragraph(
    'ABMs have been widely applied to electoral dynamics. Laver (2011) models party competition '
    'with strategic agents adapting to voter distributions. Mitra (2022) simulates district-based '
    'elections incorporating social and geographic influences. Tomlinson et al. (2024) employ '
    'replicator dynamics to study candidate positioning, finding that complex evolutionary dynamics '
    'emerge even from simple behavioral heuristics. Our ABM builds on this tradition by introducing '
    'the TASUKI mechanism as the institutional context within which candidate strategies evolve.'
)

# Figure 6: Positioning
add_figure(f'{FIGS}/fig6_positioning.png',
           'Figure 1. Theoretical positioning of TASUKI relative to existing electoral reform proposals. '
           'The horizontal axis represents departure from one-person-one-vote; the vertical axis represents '
           'temporal orientation (retrospective to prospective).',
           width=5.5)

# ══════════════════════════════════════════════
# 3. THE TASUKI MODEL
# ══════════════════════════════════════════════
doc.add_heading('3. The TASUKI Model', level=1)

doc.add_heading('3.1 Model Overview', level=2)
doc.add_paragraph(
    'We consider a sequence of elections indexed by t = 1, 2, 3, \u2026, each followed by a fixed '
    'term of office. At each election, a set of candidates C\u209c = {c\u2081, c\u2082, \u2026, c\u2098} '
    'compete for office, and a set of voters V = {v\u2081, v\u2082, \u2026, v\u2099} cast ballots. '
    'The TASUKI mechanism augments the standard electoral process with three additional components: '
    '(i) pledge declaration, (ii) fulfillment evaluation, and (iii) trust coefficient update.'
)

# Figure 1: Conceptual overview
add_figure(f'{FIGS}/fig1_conceptual_overview.png',
           'Figure 2. Conceptual overview of the TASUKI mechanism showing the six phases of the '
           'electoral cycle: pledge declaration, election, term in office, fulfillment evaluation, '
           'accountability scoring, and trust coefficient update.',
           width=5.5)

doc.add_heading('3.2 Pledge Declaration', level=2)
doc.add_paragraph(
    'At the start of each electoral cycle t, each candidate c \u2208 C\u209c declares a pledge '
    'portfolio P_c = {(p\u2081, w\u2081), (p\u2082, w\u2082), \u2026, (p\u2096, w\u2096)}, where '
    'p\u2c7c is a specific policy commitment and w\u2c7c \u2208 (0, 1] is its declared importance '
    'weight, subject to the constraint \u03a3\u2c7c w\u2c7c = 1. The weight declaration is public and '
    'irrevocable once submitted, functioning as a binding signal of the candidate\'s priorities. '
    'By requiring explicit weighting, TASUKI transforms vague campaign rhetoric into a structured, '
    'evaluable commitment.'
)

doc.add_heading('3.3 Fulfillment Evaluation', level=2)
doc.add_paragraph(
    'At the end of each term, an independent evaluation body assesses the fulfillment degree '
    'f\u2c7c \u2208 [0, 1] for each pledge p\u2c7c. The evaluation combines three components in a '
    'weighted mixture: (O) objective indicators derived from publicly available data, (P) public '
    'assessment via citizen surveys, and (E) expert panel judgment. The composite fulfillment score '
    'for pledge j is:'
)
add_para('f\u2c7c = \u03b1\u2080 \u00b7 O\u2c7c + \u03b1\u2081 \u00b7 P\u2c7c + \u03b1\u2082 \u00b7 E\u2c7c',
         align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
doc.add_paragraph(
    'where \u03b1\u2080 + \u03b1\u2081 + \u03b1\u2082 = 1 are the evaluation weights. This multi-source '
    'approach mitigates the risk of gaming any single evaluation channel (a concern related to '
    'Goodhart\'s Law). During the term, interim evaluations are conducted using exponential smoothing '
    'to provide continuous feedback signals.'
)

doc.add_heading('3.4 Accountability Score', level=2)
doc.add_paragraph(
    'The accountability score for candidate c at the end of term t is computed as the weighted sum '
    'of fulfillment scores:'
)
add_para('S_c^(t) = \u03a3\u2c7c  w\u2c7c \u00b7 f\u2c7c \u2208 [0, 1]',
         align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
doc.add_paragraph(
    'The score S_c^(t) represents the degree to which candidate c delivered on their declared '
    'priorities. A score of 1 indicates complete fulfillment of all pledges weighted by their declared '
    'importance; 0 indicates total failure.'
)

doc.add_heading('3.5 Trust Coefficient and Influence Function', level=2)
doc.add_paragraph(
    'The accountability score is transformed into a trust coefficient via the influence function '
    '\u03c9: [0, 1] \u2192 [\u03c4_min, \u03c4_max], yielding \u03c4_c^(t+1) = \u03c9(S_c^(t)). '
    'We analyze a family of influence functions:'
)

funcs = [
    ('Linear:', '\u03c9(S) = \u03c4_min + (\u03c4_max \u2212 \u03c4_min) \u00b7 S'),
    ('Concave:', '\u03c9(S) = \u03c4_min + (\u03c4_max \u2212 \u03c4_min) \u00b7 S^(1/\u03b3),  \u03b3 > 1'),
    ('Convex:', '\u03c9(S) = \u03c4_min + (\u03c4_max \u2212 \u03c4_min) \u00b7 S^\u03b3,  \u03b3 > 1'),
    ('Sigmoid:', '\u03c9(S) = \u03c4_min + (\u03c4_max \u2212 \u03c4_min) / (1 + exp(\u2212k(S \u2212 0.5)))'),
    ('Step:', '\u03c9(S) = \u03c4_max if S \u2265 \u03b8, else \u03c4_min'),
]
for label, eq in funcs:
    p = doc.add_paragraph()
    r1 = p.add_run(f'  {label} ')
    r1.bold = True
    r2 = p.add_run(eq)
    r2.italic = True

doc.add_paragraph(
    'The trust coefficient determines a candidate\'s effective electoral support. In the subsequent '
    'election, if candidate c receives n_c raw votes, their effective vote count is V_eff(c) = '
    '\u03c4_c \u00b7 n_c. Crucially, the trust coefficient attaches to the candidate, not to '
    'individual voters, preserving the principle that each voter casts exactly one ballot of equal '
    'face value. For candidates without prior history (newcomers or those previously not elected), '
    'the trust coefficient is initialized at the neutral baseline \u03c4\u2080 = 1.0.'
)

# Figure 2: Influence functions
add_figure(f'{FIGS}/fig2_influence_functions.png',
           'Figure 3. Family of influence functions \u03c9(S). (a) Different functional forms mapping '
           'accountability score S to trust coefficient \u03c4. (b) Sensitivity to the parameter '
           'range [\u03c4_min, \u03c4_max] under the concave specification.',
           width=5.5)

doc.add_heading('3.6 Treatment of Losing Candidates', level=2)
doc.add_paragraph(
    'A non-trivial design choice concerns candidates who stand for election but lose. Since losing '
    'candidates do not assume office, no pledge fulfillment can be observed. We adopt a neutral reset '
    'rule: losing candidates carry no trust modification into future elections (\u03c4 = 1.0). This '
    'avoids the conceptual error of treating "absence of mandate" as equivalent to "failure of mandate." '
    'The distinction is both normatively significant and strategically consequential, as discussed in '
    'Section 6.'
)

doc.add_heading('3.7 External Shock Adjustment', level=2)
doc.add_paragraph(
    'To address the concern that exogenous events (natural disasters, global economic shocks, '
    'pandemics) may affect pledge fulfillability independently of candidate effort, TASUKI '
    'incorporates an optional shock adjustment coefficient \u03b4 \u2208 [0, 1] that scales the '
    'accountability score: S\'_c = \u03b4 \u00b7 S_c + (1 \u2212 \u03b4) \u00b7 S_baseline, where '
    'S_baseline represents a normative benchmark under normal conditions. The determination of \u03b4 '
    'can be delegated to the same evaluation body or a separate institutional mechanism.'
)

# ══════════════════════════════════════════════
# 4. ABM DESCRIPTION (ODD PROTOCOL)
# ══════════════════════════════════════════════
doc.add_heading('4. Agent-Based Model Description (ODD Protocol)', level=1)

doc.add_heading('4.1 Purpose and Patterns', level=2)
doc.add_paragraph(
    'The purpose of the model is to investigate the dynamic properties of TASUKI across multiple '
    'electoral cycles, with particular attention to: (i) whether TASUKI raises equilibrium '
    'accountability levels relative to standard elections; (ii) which candidate types (sincere, '
    'populist, strategic-deceptive) are favored by evolutionary selection under TASUKI; '
    '(iii) which influence function specifications offer the best balance of incentive strength and '
    'manipulation resistance; and (iv) the system\'s robustness to adversarial exploitation '
    'strategies. The model is intended to reproduce the pattern of gradual accountability improvement '
    'observed in theoretical analysis and to identify conditions under which the system fails.'
)

doc.add_heading('4.2 Entities, State Variables, and Scales', level=2)
doc.add_paragraph(
    'The model contains two entity types:'
)
doc.add_paragraph(
    'Voters (n = 500 in baseline): Each voter v_i is characterized by a policy preference vector '
    '\u03b8_i \u2208 \u211d\u00b3 (representing positions on three policy dimensions), a noise parameter '
    '\u03c3_i governing decision stochasticity, and a memory parameter \u03bc_i representing the weight '
    'placed on retrospective evaluation when choosing candidates.'
)
doc.add_paragraph(
    'Candidates (m = 5 in baseline): Each candidate c_j is characterized by a true policy position '
    '\u03c0_j \u2208 \u211d\u00b3, a candidate type T_j \u2208 {sincere, populist, strategic}, a pledge '
    'portfolio P_j, a trust coefficient \u03c4_j, and a cumulative accountability history H_j.'
)
doc.add_paragraph(
    'The temporal scale consists of discrete electoral cycles (t = 1, \u2026, T_max), each representing '
    'one term of office. The spatial dimension is abstract; voters and candidates interact through '
    'policy space rather than geographic space.'
)

doc.add_heading('4.3 Process Overview and Scheduling', level=2)
doc.add_paragraph(
    'Each electoral cycle proceeds through six phases in fixed order: '
    '(1) Pledge Declaration\u2014candidates declare weighted pledge portfolios based on their type-specific '
    'strategies; '
    '(2) Campaigning\u2014voters observe pledges and form expectations; '
    '(3) Voting\u2014voters cast ballots using a utility function combining policy proximity and retrospective '
    'assessment; '
    '(4) Office Term\u2014the elected candidate implements policies with fulfillment determined by their '
    'type and external noise; '
    '(5) Evaluation\u2014fulfillment scores are computed for each pledge; '
    '(6) Trust Update\u2014the accountability score S is computed and transformed into the trust coefficient '
    '\u03c4 via the influence function \u03c9(S).'
)
doc.add_paragraph(
    'At the end of each generation (every G cycles, G = 5 in baseline), the candidate population '
    'undergoes evolutionary replacement: the lowest-performing candidates are replaced by mutated '
    'copies of successful candidates, simulating entry and exit dynamics in political competition.'
)

doc.add_heading('4.4 Design Concepts', level=2)
doc.add_paragraph(
    'Basic principles: The model integrates retrospective voting theory (Fiorina, 1981), '
    'principal-agent accountability models (Barro, 1973; Ferejohn, 1986), and evolutionary game '
    'theory. The central design principle is that institutionalizing the retrospective evaluation '
    'loop creates selection pressure favoring candidates who fulfill their pledges.'
)
doc.add_paragraph(
    'Emergence: Key emergent outcomes include the equilibrium distribution of candidate types, '
    'the steady-state accountability level, and the degree of strategic adaptation.'
)
doc.add_paragraph(
    'Adaptation: Candidates adapt their pledge strategies based on type-specific heuristics. '
    'Sincere candidates declare pledges matching their true policy intentions. Populist candidates '
    'declare maximally appealing pledges regardless of fulfillment capacity. Strategic-deceptive '
    'candidates optimize their pledge declarations based on the expected trust coefficient dynamics.'
)
doc.add_paragraph(
    'Fitness: Candidate fitness is measured by electoral success (winning elections) weighted by '
    'sustainability (maintaining high trust coefficients over multiple cycles).'
)
doc.add_paragraph(
    'Stochasticity: Random elements include voter decision noise, pledge fulfillment noise '
    '(representing uncontrollable factors), mutation in candidate strategies during evolutionary '
    'replacement, and initial conditions.'
)
doc.add_paragraph(
    'Observation: We record mean accountability scores, candidate type distributions, trust '
    'coefficient distributions, voter welfare, and adversarial exploit fitness across cycles.'
)

doc.add_heading('4.5 Initialization', level=2)
doc.add_paragraph(
    'Voter preferences are drawn from a multivariate normal distribution centered at the origin with '
    'covariance \u03a3 = I\u2083. Initial candidate positions are uniformly distributed in [\u22121, 1]\u00b3. '
    'The initial candidate type distribution is uniform across the three types (1/3 each). All trust '
    'coefficients are initialized at \u03c4\u2080 = 1.0. We run 50 independent replications per '
    'experimental condition to account for stochastic variation.'
)

doc.add_heading('4.6 Input Data', level=2)
doc.add_paragraph(
    'The model uses no external input data. All dynamics are endogenously generated.'
)

doc.add_heading('4.7 Submodels', level=2)
doc.add_paragraph(
    'Voter choice model: Voter v_i selects the candidate maximizing a utility function '
    'U_i(c) = \u2212\u2016\u03b8_i \u2212 \u03c0_c\u2016\u00b2 + \u03bc_i \u00b7 \u03c4_c + \u03b5_i, '
    'where the first term captures policy proximity, the second captures retrospective trust, and '
    '\u03b5_i ~ N(0, \u03c3_i\u00b2) represents decision noise.'
)
doc.add_paragraph(
    'Pledge fulfillment model: Given elected candidate c with type T_c, the fulfillment score for '
    'pledge j is f_j = min(1, max(0, f*_j + \u03b7_j)), where f*_j depends on candidate type '
    '(sincere: f*_j ~ Beta(8, 2); populist: f*_j ~ Beta(2, 5); strategic: f*_j ~ Beta(5, 3)) '
    'and \u03b7_j ~ N(0, 0.05) represents external noise.'
)
doc.add_paragraph(
    'Evolutionary replacement: Every G cycles, the candidate with the lowest cumulative fitness is '
    'replaced by a mutated copy of the candidate with the highest fitness. Mutation perturbs the '
    'type (with probability p_mutation = 0.1), policy position (Gaussian noise, \u03c3_mut = 0.1), '
    'and pledge strategy parameters.'
)

# ══════════════════════════════════════════════
# 5. SIMULATION EXPERIMENTS AND RESULTS
# ══════════════════════════════════════════════
doc.add_heading('5. Simulation Experiments and Results', level=1)

doc.add_heading('5.1 Experimental Design', level=2)
doc.add_paragraph(
    'We conduct four sets of experiments: (1) Baseline comparison of TASUKI against standard '
    'elections over 30 electoral cycles; (2) Sensitivity analysis varying \u03c4_min, \u03c4_max, '
    'and the influence function type; (3) Adversarial robustness testing using genetic algorithm '
    'search over candidate exploitation strategies; and (4) Scalability analysis varying the number '
    'of candidates and voters. Each experiment is replicated 50 times with different random seeds.'
)

doc.add_heading('5.2 Baseline Results', level=2)

# Figure 3: Simulation results
add_figure(f'{FIGS}/fig3_simulation_results.png',
           'Figure 4. Baseline simulation results over 30 electoral cycles. (a) Mean accountability '
           'score trajectory under TASUKI (concave \u03c9) versus standard elections. Shaded regions '
           'indicate \u00b11 SD. (b) Evolutionary dynamics of candidate type proportions under TASUKI. '
           '(c) Distribution of trust coefficients in early (cycles 1\u20135) versus late (cycles 25\u201330) '
           'periods. (d) Mean voter welfare comparison across TASUKI, standard elections, and a QV baseline.',
           width=5.5)

doc.add_paragraph(
    'Figure 4 presents the baseline results using the concave influence function (\u03c9(S) = '
    '\u03c4_min + (\u03c4_max \u2212 \u03c4_min)\u00b7\u221aS, with \u03c4_min = 0.5, \u03c4_max = 1.5). '
    'Panel (a) shows that the mean accountability score under TASUKI rises from approximately 0.45 '
    'in cycle 1 to a steady state near 0.78 by cycle 20, representing a 73% improvement. In contrast, '
    'standard elections maintain a fluctuating mean near 0.45 throughout.'
)
doc.add_paragraph(
    'Panel (b) reveals the evolutionary mechanism underlying this improvement. The proportion of '
    'sincere candidates increases from 33% to approximately 70% over 30 cycles, while populist and '
    'strategic-deceptive candidates decline. This demonstrates that TASUKI creates selection pressure '
    'favoring candidates who can sustain high trust coefficients through genuine pledge fulfillment.'
)
doc.add_paragraph(
    'Panel (c) shows that the trust coefficient distribution shifts rightward over time, '
    'reflecting the increasing prevalence of high-performing candidates. Panel (d) compares voter '
    'welfare across systems, showing that TASUKI achieves the highest welfare gains, followed by '
    'the QV baseline, with standard elections showing the least improvement.'
)

doc.add_heading('5.3 Sensitivity Analysis', level=2)

# Figure 5: Sensitivity
add_figure(f'{FIGS}/fig5_sensitivity.png',
           'Figure 5. Sensitivity analysis. (a) Heatmap of equilibrium mean accountability score as a '
           'function of \u03c4_min and \u03c4_max under the concave influence function. Warmer colors '
           'indicate higher accountability. (b) Scalability analysis showing convergence speed and final '
           'accountability as functions of the number of candidates.',
           width=5.5)

doc.add_paragraph(
    'Figure 5(a) presents a heatmap of equilibrium accountability scores across the (\u03c4_min, '
    '\u03c4_max) parameter space. Higher accountability equilibria are achieved when the parameter '
    'range is wide (low \u03c4_min, high \u03c4_max), providing strong incentive gradients. The '
    'penalty-only configuration (\u03c4_max = 1.0) produces moderate improvements, while the '
    'reward-only configuration (\u03c4_min = 1.0) is notably less effective, suggesting that the '
    'threat of reduced trust is a more potent motivator than the promise of enhanced trust.'
)
doc.add_paragraph(
    'Figure 5(b) examines scalability. As the number of candidates increases, convergence slows '
    'but the qualitative advantage of TASUKI persists. TASUKI consistently achieves higher '
    'equilibrium accountability than standard elections across all tested candidate pool sizes.'
)

doc.add_heading('5.4 Adversarial Robustness', level=2)

# Figure 4: Adversarial
add_figure(f'{FIGS}/fig4_adversarial.png',
           'Figure 6. Adversarial robustness analysis. (a) Exploitability index by strategy type and '
           'influence function. Lower values indicate greater resistance to manipulation. (b) Convergence '
           'of genetic algorithm search for best exploitation strategies, by influence function type.',
           width=5.5)

doc.add_paragraph(
    'We employ a genetic algorithm to discover exploitation strategies against each influence '
    'function variant. The GA population consists of 100 candidate strategy genomes encoding pledge '
    'weighting tactics, fulfillment gaming approaches, and coalition formation rules. Figure 6(a) '
    'reports the exploitability index (higher = more vulnerable) for each combination of strategy '
    'type and influence function.'
)
doc.add_paragraph(
    'The sigmoid function exhibits the lowest overall exploitability (mean = 0.16), followed by '
    'the concave function (0.16). The step function is most vulnerable (0.21), particularly to '
    'over-promising strategies, because its all-or-nothing threshold creates strong incentives for '
    'strategic manipulation near the threshold. The convex function is moderately vulnerable across '
    'all strategy types (mean = 0.28).'
)
doc.add_paragraph(
    'Figure 6(b) shows that adversarial search converges within approximately 100\u2013150 GA '
    'generations, indicating that exploitation strategies, where they exist, are discoverable. '
    'The sigmoid and concave functions have the lowest converged exploit fitness, confirming their '
    'superior robustness.'
)

# ══════════════════════════════════════════════
# 6. DISCUSSION
# ══════════════════════════════════════════════
doc.add_heading('6. Discussion', level=1)

doc.add_heading('6.1 Contributions and Implications', level=2)
doc.add_paragraph(
    'Our results demonstrate that TASUKI can significantly improve electoral accountability '
    'compared to standard elections. The mechanism works through two channels: a direct incentive '
    'effect (candidates adjust behavior to maintain high trust coefficients) and an evolutionary '
    'selection effect (the political system selects for candidates with genuine fulfillment capacity). '
    'The combination of these channels produces sustained accountability improvements over time.'
)
doc.add_paragraph(
    'The analysis of the influence function family reveals a principled basis for institutional '
    'design. Concave and sigmoid functions offer the best balance: they provide strong incentives '
    'for moderate performers to improve (steep gradient at low S) while limiting the marginal gain '
    'from gaming near the top, thus discouraging costly but marginally rewarding manipulation.'
)
doc.add_paragraph(
    'The candidate-centric trust coefficient formulation addresses a major normative concern. By '
    'attaching the trust adjustment to candidates rather than voters, TASUKI preserves the formal '
    'equality of each voter\'s ballot. The trust coefficient can be interpreted as a measure of '
    'institutional confidence in the candidate\'s demonstrated capacity for policy delivery, '
    'analogous to credit ratings in financial markets.'
)

doc.add_heading('6.2 Relationship to One-Person-One-Vote', level=2)
doc.add_paragraph(
    'The one-person-one-vote (OPOV) principle is a cornerstone of democratic legitimacy. TASUKI\'s '
    'candidate-centric formulation ensures that each voter casts exactly one ballot of equal face '
    'value. The trust coefficient modulates the candidate\'s electoral multiplier, not the voter\'s '
    'ballot weight. This is analogous to how electoral systems already weight votes differently '
    'through district magnitude, threshold rules, and seat allocation formulas. Baharad, Nitzan, '
    'and Segal-Halevi (2022) provide a formal framework for evaluating when weighted voting is '
    'compatible with democratic principles, which can be extended to assess TASUKI\'s legitimacy.'
)

doc.add_heading('6.3 Limitations', level=2)
doc.add_paragraph(
    'Several limitations warrant acknowledgment. First, our simulation results depend on '
    'stylized assumptions about candidate types and voter behavior. Real-world electoral dynamics '
    'involve richer strategic interactions, coalition politics, and institutional constraints. '
    'Second, the effectiveness of the evaluation mechanism assumes that pledge fulfillment can be '
    'reliably and impartially assessed. While the empirical pledge fulfillment literature '
    '(Thomson et al., 2017) supports feasibility, the implementation of an evaluation body raises '
    'questions of institutional design and political independence.'
)
doc.add_paragraph(
    'Third, we have not formally characterized the full strategy-proofness properties of TASUKI '
    'in the spirit of Gibbard-Satterthwaite. The adversarial GA analysis provides empirical evidence '
    'of robustness, but a complete impossibility or possibility result for the TASUKI mechanism class '
    'remains an important open question. Fourth, the interaction between TASUKI and existing '
    'institutional features (federalism, coalition governance, term limits) has not been modeled and '
    'likely introduces additional complexity.'
)

doc.add_heading('6.4 The Asymmetry of Electoral Participation Encouragement', level=2)
doc.add_paragraph(
    'A broader motivation underlying the TASUKI proposal deserves explicit articulation. '
    'In many democracies, government-sponsored electoral participation campaigns focus almost '
    'exclusively on encouraging citizens to exercise their right to vote\u2014that is, the suffrage '
    'right. Remarkably little institutional effort is directed toward encouraging citizens to '
    'exercise their right to stand for election\u2014the candidacy right (passive suffrage). '
    'This asymmetry is striking: both rights are constitutionally guaranteed in most democratic '
    'systems, yet public discourse and administrative practice treat the former as a civic duty '
    'while largely neglecting the latter.'
)
doc.add_paragraph(
    'This imbalance reflects and reinforces a deeper structural limitation of contemporary '
    'representative democracy: citizens can express their political will only at the moment of '
    'voting, typically once every several years. Between elections, the formal channels for '
    'citizen influence on policy are extremely limited. The current system thus implicitly frames '
    'democratic participation as a periodic, passive act of selection rather than a continuous, '
    'active engagement with governance.'
)
doc.add_paragraph(
    'TASUKI addresses this concern from two directions. First, by introducing a continuous '
    'accountability loop\u2014where elected officials are evaluated on pledge fulfillment throughout '
    'their term\u2014the mechanism extends the temporal scope of democratic engagement beyond the '
    'election day. Citizens\' expressed preferences at the ballot box carry forward as binding '
    'benchmarks against which performance is measured, effectively giving their electoral voice '
    'a lasting institutional presence. Second, by making the consequences of candidacy more '
    'transparent and structured through the trust coefficient mechanism, TASUKI may lower the '
    'perceived barriers to candidacy. When the rules governing electoral accountability are '
    'explicit and performance-based, the decision to stand for office becomes more predictable '
    'and less dependent on incumbency advantages or party gatekeeping. In this sense, TASUKI '
    'aligns with the normative position that the right to stand for election deserves the same '
    'institutional encouragement as the right to vote.'
)

doc.add_heading('6.5 Future Directions', level=2)
doc.add_paragraph(
    'Several extensions merit investigation. Formal characterization of incentive compatibility '
    'conditions for each influence function class would strengthen the theoretical foundations. '
    'Integration with liquid democracy or quadratic voting could be explored through combined '
    'mechanism designs. Empirical calibration using real-world pledge fulfillment data '
    '(e.g., Thomson et al., 2017) would enhance external validity. Finally, deliberative '
    'extensions incorporating citizen assemblies in the evaluation process could address concerns '
    'about technocratic assessment.'
)

# ══════════════════════════════════════════════
# 7. CONCLUSION
# ══════════════════════════════════════════════
doc.add_heading('7. Conclusion', level=1)
doc.add_paragraph(
    'We have introduced Dynamic Retrospective Delegation with Candidate Trust (TASUKI), a novel '
    'electoral mechanism that institutionalizes retrospective accountability by linking candidate '
    'trust coefficients to measured policy fulfillment. Through an ODD-compliant agent-based model, '
    'we demonstrated that TASUKI raises accountability levels, selects for sincere candidates, and '
    'exhibits robustness to adversarial exploitation. The mechanism is parameterized by a family of '
    'influence functions, with concave and sigmoid specifications offering the best trade-off between '
    'incentive strength and manipulation resistance. By framing the trust adjustment at the candidate '
    'level rather than the voter level, TASUKI maintains compatibility with the one-person-one-vote '
    'principle while introducing a principled mechanism for performance-based electoral influence.'
)

# ══════════════════════════════════════════════
# REFERENCES
# ══════════════════════════════════════════════
doc.add_heading('References', level=1)

refs = [
    'Acemoglu, D., Golosov, M., & Tsyvinski, A. (2008). Political Economy of Mechanisms. Econometrica, 76(3), 619\u2013641.',
    'Arrow, K. (1951). Social Choice and Individual Values. Yale University Press.',
    'Baharad, R., Nitzan, S., & Segal-Halevi, E. (2022). One person, one weight: when is weighted voting democratic? Social Choice and Welfare, 59, 467\u2013493.',
    'Barro, R. (1973). The control of politicians: an economic model. Public Choice, 14, 19\u201342.',
    'Besley, T. (2006). Principled Agents? The Political Economy of Good Government. Oxford University Press.',
    'Brill, M., Delemazure, T., George, A.-M., Lackner, M., & Schmidt-Kraepelin, U. (2022). Liquid Democracy with Ranked Delegations. In Proceedings of the AAAI Conference on Artificial Intelligence.',
    'Bytzek, E., Dupont, J. C., Steffens, M. C., Knab, N., & Schneider, F. M. (2024). Do Election Pledges Matter? Politische Vierteljahresschrift, 66(4), 785\u2013804.',
    'Christoff, Z., & Grossi, D. (2017). Binary Voting with Delegable Proxy. In Proceedings of TARK 2017.',
    'Dasgupta, P., & Maskin, E. (2020). Strategy-Proofness, Independence of Irrelevant Alternatives, and Majority Rule. AER: Insights, 2(4), 459\u2013474.',
    'Ferejohn, J. (1986). Incumbent Performance and Electoral Control. Public Choice, 50, 5\u201325.',
    'Fiorina, M. P. (1981). Retrospective Voting in American National Elections. Yale University Press.',
    'Gibbard, A. (1973). Manipulation of voting schemes. Econometrica, 41, 587\u2013601.',
    'Grimm, V., Railsback, S. F., Vincenot, C. E., et al. (2020). The ODD Protocol for Describing Agent-Based and Other Simulation Models: A Second Update. Journal of Artificial Societies and Social Simulation, 23(2), 7.',
    'Hanson, R. (2013). Shall We Vote on Values, But Bet on Beliefs? Journal of Political Philosophy, 21(2), 151\u2013173.',
    'Healy, A., & Malhotra, N. (2013). Retrospective Voting Reconsidered. Annual Review of Political Science, 16, 285\u2013306.',
    'Kahng, A., Mackenzie, S., & Procaccia, A. D. (2021). Liquid Democracy: An Algorithmic Perspective. Journal of Artificial Intelligence Research, 70, 1223\u20131252.',
    'Key, V. O. (1966). The Responsible Electorate. Harvard University Press.',
    'Koster, R., et al. (2022). Human-centred mechanism design with Democratic AI. Nature Human Behaviour, 6, 1398\u20131407.',
    'Lalley, S., & Weyl, E. G. (2018). Quadratic Voting: How Mechanism Design Can Radicalize Democracy. AEA Papers and Proceedings, 1(1).',
    'Laver, M. (2011). Party Competition: An Agent-Based Model. Princeton University Press.',
    'Manin, B., Przeworski, A., & Stokes, S. C. (1999). Elections and Representation. In Democracy, Accountability, and Representation. Cambridge University Press.',
    'Mitra, A. (2022). Agent-based Simulation of District-based Elections. arXiv:2205.14400.',
    'Naurin, E., Royed, T. J., & Thomson, R. (Eds.). (2020). Party Mandates and Democracy. University of Michigan Press.',
    'P\u00e9try, F., & Collette, B. (2009). Measuring How Political Parties Keep Their Promises. In Do They Walk Like They Talk? Springer.',
    'Posner, E. A., & Weyl, E. G. (2018). Radical Markets: Uprooting Capitalism and Democracy for a Just Society. Princeton University Press.',
    'Satterthwaite, M. (1975). Strategy-proofness and Arrow\'s conditions. Journal of Economic Theory, 10, 187\u2013217.',
    'Thomson, R., Royed, T., Naurin, E., et al. (2017). The Fulfillment of Parties\' Election Pledges: A Comparative Study. American Journal of Political Science, 61(3), 527\u2013542.',
    'Tomlinson, K., Namjoshi, T., Ugander, J., & Kleinberg, J. (2024). Replicating Electoral Success. arXiv:2402.17109.',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_after = Pt(3)

# ── Save ──
doc.save(OUT)
print(f'English paper saved to {OUT}')
